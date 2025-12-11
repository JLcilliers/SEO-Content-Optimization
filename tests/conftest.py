"""
Pytest fixtures and configuration for SEO Content Optimizer tests.
"""

import pytest
from pathlib import Path
import tempfile
import os

from docx import Document


@pytest.fixture
def sample_keywords_csv(tmp_path: Path) -> Path:
    """Create a sample keywords CSV file."""
    csv_path = tmp_path / "keywords.csv"
    csv_content = """keyword,search_volume,difficulty,intent
PTO insurance,1200,45,transactional
professional liability insurance,800,50,informational
how much does PTO insurance cost,150,30,informational
PTO coverage for consultants,200,35,transactional
liability protection,500,40,mixed
what is PTO insurance,300,25,informational
"""
    csv_path.write_text(csv_content)
    return csv_path


@pytest.fixture
def sample_keywords_excel(tmp_path: Path) -> Path:
    """Create a sample keywords Excel file."""
    import pandas as pd

    xlsx_path = tmp_path / "keywords.xlsx"
    data = {
        "keyword": ["PTO insurance", "liability coverage", "insurance quote"],
        "search_volume": [1000, 500, 300],
        "difficulty": [40, 35, 25],
        "intent": ["transactional", "informational", "transactional"],
    }
    df = pd.DataFrame(data)
    df.to_excel(xlsx_path, index=False)
    return xlsx_path


@pytest.fixture
def sample_docx(tmp_path: Path) -> Path:
    """Create a sample Word document."""
    docx_path = tmp_path / "sample.docx"
    doc = Document()

    doc.add_heading("Professional Liability Insurance Guide", level=1)
    doc.add_paragraph(
        "This guide covers everything you need to know about professional liability "
        "insurance. We will explain the key concepts and help you understand your coverage options."
    )
    doc.add_heading("What is Professional Liability Insurance?", level=2)
    doc.add_paragraph(
        "Professional liability insurance protects businesses and professionals from claims "
        "of negligence, errors, or omissions in their professional services. This type of "
        "coverage is essential for consultants, contractors, and service providers."
    )
    doc.add_heading("Who Needs This Coverage?", level=2)
    doc.add_paragraph(
        "Any professional who provides advice or services should consider this coverage. "
        "This includes consultants, accountants, lawyers, and healthcare providers."
    )

    doc.save(str(docx_path))
    return docx_path


@pytest.fixture
def sample_html_content() -> str:
    """Sample HTML content for testing URL parsing."""
    return """
<!DOCTYPE html>
<html>
<head>
    <title>Professional Liability Insurance | Expert Guide</title>
    <meta name="description" content="Learn about professional liability insurance coverage options for your business.">
</head>
<body>
    <header>
        <nav>Navigation content</nav>
    </header>
    <main>
        <h1>Understanding Professional Liability Insurance</h1>
        <p>Professional liability insurance is a crucial form of coverage for businesses
        that provide professional services. This guide explains everything you need to know.</p>
        <h2>Key Benefits</h2>
        <p>The main benefits include protection against claims of negligence, coverage for
        legal defense costs, and peace of mind for you and your clients.</p>
        <h2>Getting Started</h2>
        <p>Contact us today to get a quote for your professional liability insurance needs.
        Our experts are ready to help you find the right coverage.</p>
    </main>
    <footer>Footer content</footer>
</body>
</html>
"""


@pytest.fixture
def mock_llm_response():
    """Mock LLM response for testing without API calls."""
    def _mock(content: str) -> str:
        # Simple mock that adds markers around certain patterns
        result = content
        # Add markers to simulate optimization
        if "insurance" in content.lower():
            result = result.replace(
                "insurance",
                "[[[ADD]]]comprehensive insurance[[[ENDADD]]]",
                1
            )
        return result
    return _mock
