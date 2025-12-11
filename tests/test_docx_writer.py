"""Tests for Word document writing with green highlights."""

import pytest
from pathlib import Path

from docx import Document
from docx.enum.text import WD_COLOR_INDEX

from seo_content_optimizer.docx_writer import (
    DocxWriter,
    add_marked_text,
    create_simple_docx_with_highlights,
    sanitize_for_xml,
    write_optimization_result,
    MARK_START,
    MARK_END,
)
from seo_content_optimizer.models import (
    FAQItem,
    HeadingLevel,
    MetaElement,
    OptimizationResult,
    ParagraphBlock,
)


class TestDocxWriter:
    """Tests for the DocxWriter class."""

    def test_parse_marker_segments_no_markers(self):
        """Test parsing text without markers."""
        writer = DocxWriter()
        segments = writer._parse_marker_segments("Plain text without markers.")

        assert len(segments) == 1
        assert segments[0] == ("Plain text without markers.", False)

    def test_parse_marker_segments_with_markers(self):
        """Test parsing text with ADD markers."""
        writer = DocxWriter()
        text = "Original [[[ADD]]]new content[[[ENDADD]]] more original."
        segments = writer._parse_marker_segments(text)

        assert len(segments) == 3
        assert segments[0] == ("Original ", False)
        assert segments[1] == ("new content", True)
        assert segments[2] == (" more original.", False)

    def test_parse_marker_segments_multiple_markers(self):
        """Test parsing text with multiple marker sections."""
        writer = DocxWriter()
        text = "[[[ADD]]]First[[[ENDADD]]] middle [[[ADD]]]second[[[ENDADD]]] end"
        segments = writer._parse_marker_segments(text)

        assert len(segments) == 4
        assert segments[0] == ("First", True)
        assert segments[1] == (" middle ", False)
        assert segments[2] == ("second", True)
        assert segments[3] == (" end", False)

    def test_parse_marker_segments_only_markers(self):
        """Test parsing text that is entirely marked."""
        writer = DocxWriter()
        text = "[[[ADD]]]All new content here[[[ENDADD]]]"
        segments = writer._parse_marker_segments(text)

        assert len(segments) == 1
        assert segments[0] == ("All new content here", True)

    def test_write_creates_file(self, tmp_path: Path):
        """Test that write creates a file."""
        result = OptimizationResult(
            meta_elements=[
                MetaElement(
                    element_name="Title Tag",
                    current="Old Title",
                    optimized="[[[ADD]]]New[[[ENDADD]]] Title",
                    why_changed="Added keyword",
                ),
            ],
            optimized_blocks=[
                ParagraphBlock(
                    text="This is [[[ADD]]]optimized[[[ENDADD]]] content.",
                    heading_level=HeadingLevel.BODY,
                ),
            ],
            primary_keyword="test keyword",
        )

        output_path = tmp_path / "output.docx"
        writer = DocxWriter()
        created_path = writer.write(result, output_path)

        assert created_path.exists()
        assert created_path.suffix == ".docx"

    def test_write_adds_docx_extension(self, tmp_path: Path):
        """Test that .docx extension is added if missing."""
        result = OptimizationResult(primary_keyword="test")
        output_path = tmp_path / "output"

        writer = DocxWriter()
        created_path = writer.write(result, output_path)

        assert created_path.suffix == ".docx"

    def test_write_includes_reading_guide(self, tmp_path: Path):
        """Test that the reading guide is included."""
        result = OptimizationResult(primary_keyword="test")
        output_path = tmp_path / "output.docx"

        writer = DocxWriter()
        writer.write(result, output_path)

        # Read the document and check for reading guide
        doc = Document(str(output_path))
        text = "\n".join(p.text for p in doc.paragraphs)

        assert "Reading Guide" in text or any(
            "Reading Guide" in p.style.name for p in doc.paragraphs if p.style
        )

    def test_write_includes_meta_table(self, tmp_path: Path):
        """Test that meta elements table is included."""
        result = OptimizationResult(
            meta_elements=[
                MetaElement(
                    element_name="Title Tag",
                    current="Old",
                    optimized="New",
                    why_changed="Reason",
                ),
            ],
            primary_keyword="test",
        )
        output_path = tmp_path / "output.docx"

        writer = DocxWriter()
        writer.write(result, output_path)

        doc = Document(str(output_path))
        assert len(doc.tables) >= 1

    def test_write_includes_faq(self, tmp_path: Path):
        """Test that FAQ section is included when present."""
        result = OptimizationResult(
            faq_items=[
                FAQItem(
                    question="[[[ADD]]]What is SEO?[[[ENDADD]]]",
                    answer="[[[ADD]]]SEO is search engine optimization.[[[ENDADD]]]",
                ),
            ],
            primary_keyword="test",
        )
        output_path = tmp_path / "output.docx"

        writer = DocxWriter()
        writer.write(result, output_path)

        doc = Document(str(output_path))
        text = "\n".join(p.text for p in doc.paragraphs)

        assert "Frequently Asked Questions" in text or "FAQ" in text


class TestCreateSimpleDocx:
    """Tests for the simple docx creation helper."""

    def test_create_simple_docx(self, tmp_path: Path):
        """Test creating a simple docx with highlights."""
        text = "Normal text [[[ADD]]]highlighted text[[[ENDADD]]] more normal."
        output_path = tmp_path / "simple.docx"

        created_path = create_simple_docx_with_highlights(text, output_path)

        assert created_path.exists()
        doc = Document(str(created_path))
        assert len(doc.paragraphs) > 0


class TestWriteOptimizationResult:
    """Tests for the convenience function."""

    def test_write_optimization_result(self, tmp_path: Path):
        """Test the convenience function."""
        result = OptimizationResult(
            meta_elements=[],
            optimized_blocks=[
                ParagraphBlock(text="Test content", heading_level=HeadingLevel.BODY),
            ],
            primary_keyword="test",
        )
        output_path = tmp_path / "result.docx"

        created_path = write_optimization_result(result, output_path)

        assert created_path.exists()


class TestAddMarkedText:
    """Tests for the add_marked_text function."""

    def test_add_marked_text_no_markers(self, tmp_path: Path):
        """Test adding text without markers."""
        doc = Document()
        para = doc.add_paragraph()
        add_marked_text(para, "Plain text without markers.")

        assert len(para.runs) == 1
        assert para.runs[0].text == "Plain text without markers."
        assert para.runs[0].font.highlight_color is None

    def test_add_marked_text_with_markers(self, tmp_path: Path):
        """Test adding text with markers creates highlighted runs."""
        doc = Document()
        para = doc.add_paragraph()
        add_marked_text(para, f"Normal {MARK_START}highlighted{MARK_END} text.")

        # Should have 3 runs: normal, highlighted, normal
        assert len(para.runs) == 3
        assert para.runs[0].text == "Normal "
        assert para.runs[0].font.highlight_color is None
        assert para.runs[1].text == "highlighted"
        assert para.runs[1].font.highlight_color == WD_COLOR_INDEX.BRIGHT_GREEN
        assert para.runs[2].text == " text."
        assert para.runs[2].font.highlight_color is None

    def test_add_marked_text_multiple_markers(self, tmp_path: Path):
        """Test adding text with multiple marker sections."""
        doc = Document()
        para = doc.add_paragraph()
        text = f"{MARK_START}First{MARK_END} middle {MARK_START}second{MARK_END} end"
        add_marked_text(para, text)

        assert len(para.runs) == 4
        assert para.runs[0].text == "First"
        assert para.runs[0].font.highlight_color == WD_COLOR_INDEX.BRIGHT_GREEN
        assert para.runs[1].text == " middle "
        assert para.runs[1].font.highlight_color is None
        assert para.runs[2].text == "second"
        assert para.runs[2].font.highlight_color == WD_COLOR_INDEX.BRIGHT_GREEN
        assert para.runs[3].text == " end"
        assert para.runs[3].font.highlight_color is None


class TestDocxHighlightValidation:
    """Tests that validate DOCX output has proper green highlights."""

    def test_docx_contains_bright_green_highlight(self, tmp_path: Path):
        """Test that generated DOCX has at least one BRIGHT_GREEN highlight."""
        result = OptimizationResult(
            meta_elements=[
                MetaElement(
                    element_name="Title Tag",
                    current="Old Title",
                    optimized=f"{MARK_START}New Optimized{MARK_END} Title",
                    why_changed="Added keyword",
                ),
            ],
            optimized_blocks=[
                ParagraphBlock(
                    text=f"This is {MARK_START}optimized SEO{MARK_END} content.",
                    heading_level=HeadingLevel.BODY,
                ),
            ],
            primary_keyword="SEO",
        )

        output_path = tmp_path / "highlight_test.docx"
        writer = DocxWriter()
        writer.write(result, output_path)

        # Open and verify highlights exist
        doc = Document(str(output_path))
        green_highlights_found = 0

        for para in doc.paragraphs:
            for run in para.runs:
                if run.font.highlight_color == WD_COLOR_INDEX.BRIGHT_GREEN:
                    green_highlights_found += 1

        # Also check table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            if run.font.highlight_color == WD_COLOR_INDEX.BRIGHT_GREEN:
                                green_highlights_found += 1

        assert green_highlights_found >= 2, (
            f"Expected at least 2 green highlights, found {green_highlights_found}"
        )

    def test_docx_is_valid_word_document(self, tmp_path: Path):
        """Test that generated DOCX is a valid Word document that can be opened."""
        result = OptimizationResult(
            meta_elements=[
                MetaElement(
                    element_name="Title Tag",
                    current="Current",
                    optimized=f"{MARK_START}Optimized{MARK_END}",
                    why_changed="SEO",
                ),
                MetaElement(
                    element_name="Meta Description",
                    current="Old description",
                    optimized=f"New {MARK_START}keyword-rich{MARK_END} description.",
                    why_changed="Added keywords",
                ),
            ],
            optimized_blocks=[
                ParagraphBlock(
                    text=f"{MARK_START}SEO Optimized Heading{MARK_END}",
                    heading_level=HeadingLevel.H1,
                ),
                ParagraphBlock(
                    text=f"Regular paragraph with {MARK_START}green highlighted text{MARK_END}.",
                    heading_level=HeadingLevel.BODY,
                ),
            ],
            faq_items=[
                FAQItem(
                    question=f"{MARK_START}What is SEO?{MARK_END}",
                    answer=f"{MARK_START}SEO stands for Search Engine Optimization.{MARK_END}",
                ),
            ],
            primary_keyword="SEO",
            secondary_keywords=["optimization", "search"],
        )

        output_path = tmp_path / "valid_docx_test.docx"
        writer = DocxWriter()
        created_path = writer.write(result, output_path)

        # Verify the file exists and can be opened without errors
        assert created_path.exists()
        doc = Document(str(created_path))

        # Verify basic structure
        assert len(doc.paragraphs) > 0
        assert len(doc.tables) >= 1  # Meta table should exist

        # Verify content is present
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Reading Guide" in all_text
        assert "OPTIMIZED CONTENT" in all_text
        assert "SEO Optimized Heading" in all_text

    def test_simple_docx_with_highlights_creates_valid_file(self, tmp_path: Path):
        """Test create_simple_docx_with_highlights produces valid highlighted doc."""
        text = f"Start {MARK_START}green text here{MARK_END} and {MARK_START}more green{MARK_END} end."
        output_path = tmp_path / "simple_highlight.docx"

        created_path = create_simple_docx_with_highlights(text, output_path)

        # Verify file exists
        assert created_path.exists()

        # Open and verify highlights
        doc = Document(str(created_path))
        green_count = 0

        for para in doc.paragraphs:
            for run in para.runs:
                if run.font.highlight_color == WD_COLOR_INDEX.BRIGHT_GREEN:
                    green_count += 1

        assert green_count >= 2, f"Expected at least 2 green highlights, found {green_count}"

    def test_reading_guide_has_highlighted_sample(self, tmp_path: Path):
        """Test that reading guide contains a highlighted sample."""
        result = OptimizationResult(primary_keyword="test")
        output_path = tmp_path / "reading_guide_test.docx"

        writer = DocxWriter()
        writer.write(result, output_path)

        doc = Document(str(output_path))

        # Find the paragraph with "green like this" and verify it has highlight
        found_highlighted_sample = False
        for para in doc.paragraphs:
            if "green" in para.text.lower():
                for run in para.runs:
                    if run.font.highlight_color == WD_COLOR_INDEX.BRIGHT_GREEN:
                        found_highlighted_sample = True
                        break

        assert found_highlighted_sample, "Reading guide should have a highlighted sample"


class TestSanitizeForXml:
    """Tests for XML character sanitization."""

    def test_sanitize_removes_null_bytes(self):
        """Test that null bytes are removed."""
        text = "Hello\x00World"
        result = sanitize_for_xml(text)
        assert result == "HelloWorld"

    def test_sanitize_removes_control_characters(self):
        """Test that control characters are removed."""
        # Characters 0x01-0x08, 0x0B, 0x0C, 0x0E-0x1F should be removed
        text = "Hello\x01\x02\x03\x04\x05\x06\x07\x08World"
        result = sanitize_for_xml(text)
        assert result == "HelloWorld"

    def test_sanitize_preserves_valid_whitespace(self):
        """Test that tab, newline, carriage return are preserved."""
        text = "Hello\t\n\rWorld"
        result = sanitize_for_xml(text)
        assert result == "Hello\t\n\rWorld"

    def test_sanitize_preserves_normal_text(self):
        """Test that normal text is unchanged."""
        text = "Hello World! This is normal text."
        result = sanitize_for_xml(text)
        assert result == text

    def test_sanitize_handles_empty_string(self):
        """Test handling of empty string."""
        result = sanitize_for_xml("")
        assert result == ""

    def test_sanitize_handles_none(self):
        """Test handling of None."""
        result = sanitize_for_xml(None)
        assert result is None

    def test_sanitize_removes_extended_control_chars(self):
        """Test that extended control characters (0x7F-0x9F) are removed."""
        text = "Hello\x7f\x80\x9fWorld"
        result = sanitize_for_xml(text)
        assert result == "HelloWorld"

    def test_sanitize_preserves_unicode(self):
        """Test that unicode characters are preserved."""
        text = "Héllo Wörld 中文 日本語"
        result = sanitize_for_xml(text)
        assert result == text

    def test_sanitize_with_markers_and_control_chars(self, tmp_path: Path):
        """Test that sanitization works with marker content."""
        # Create a document with text containing control characters
        text_with_control = f"Normal text {MARK_START}highlighted\x00text{MARK_END} end."
        doc = Document()
        para = doc.add_paragraph()
        add_marked_text(para, text_with_control)

        # Verify the document can be saved and opened without errors
        output_path = tmp_path / "sanitized_test.docx"
        doc.save(str(output_path))

        # Open the document to verify it's valid
        opened_doc = Document(str(output_path))
        all_text = "".join(p.text for p in opened_doc.paragraphs)
        assert "highlightedtext" in all_text  # Control char removed
        assert "\x00" not in all_text
