"""
Content extraction from various sources (URLs and Word documents).

This module handles fetching and parsing content from:
- Web URLs (using FireCrawl for structured extraction, falling back to trafilatura)
- Word documents (.docx files using python-docx)

The FireCrawl integration provides superior content extraction that preserves
the visual structure of web pages including headings, tables, and lists.
"""

import os
import re
from datetime import datetime
from pathlib import Path
from typing import Optional, Union
from urllib.parse import urlparse

import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.style import WD_STYLE_TYPE

try:
    import trafilatura
except ImportError:
    trafilatura = None  # type: ignore

from .models import (
    ContentBlock,
    DocxContent,
    HeadingLevel,
    PageMeta,
    ParagraphBlock,
    # V2 Architecture imports
    Block,
    Run,
    ContentDocument,
    BlockType,
)
from .text_repair import repair_text, repair_content_blocks, validate_text_quality


# Default headers for web requests
DEFAULT_HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
        "(KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
    ),
    "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
    "Accept-Language": "en-US,en;q=0.5",
}

# Mapping of Word heading styles to HeadingLevel
HEADING_STYLE_MAP = {
    "Heading 1": HeadingLevel.H1,
    "Heading 2": HeadingLevel.H2,
    "Heading 3": HeadingLevel.H3,
    "Heading 4": HeadingLevel.H4,
    "Heading 5": HeadingLevel.H5,
    "Heading 6": HeadingLevel.H6,
    "Title": HeadingLevel.H1,
}


class ContentExtractionError(Exception):
    """Raised when content extraction fails."""
    pass


# Extraction quality thresholds
MIN_WORD_COUNT = 100  # Minimum words for "good" extraction
MIN_BLOCK_COUNT = 3   # Minimum content blocks
MIN_HEADING_COUNT = 1  # At least one heading expected


def _decode_html_safely(response) -> str:
    """
    Decode HTTP response to string with proper encoding detection.

    CRITICAL: This prevents mojibake by ensuring correct encoding.

    Detection order:
    1. Content-Type header charset
    2. HTML meta charset tag
    3. charset_normalizer detection (if available)
    4. UTF-8 with error handling

    Args:
        response: requests.Response object

    Returns:
        Decoded HTML string
    """
    import sys

    content_bytes = response.content

    # 1. Try Content-Type header charset
    content_type = response.headers.get('Content-Type', '')
    if 'charset=' in content_type.lower():
        charset = content_type.split('charset=')[-1].split(';')[0].strip().strip('"\'')
        try:
            return content_bytes.decode(charset)
        except (UnicodeDecodeError, LookupError) as e:
            print(f"DEBUG: Header charset {charset} failed: {e}", file=sys.stderr)

    # 2. Try to find meta charset in first 8KB
    head_bytes = content_bytes[:8192]
    try:
        head_text = head_bytes.decode('ascii', errors='ignore')

        # Look for <meta charset="...">
        import re
        charset_match = re.search(r'<meta[^>]+charset=["\']?([^"\'>\s]+)', head_text, re.I)
        if charset_match:
            charset = charset_match.group(1)
            try:
                return content_bytes.decode(charset)
            except (UnicodeDecodeError, LookupError) as e:
                print(f"DEBUG: Meta charset {charset} failed: {e}", file=sys.stderr)

        # Look for <meta http-equiv="Content-Type" content="...charset=...">
        content_type_match = re.search(
            r'<meta[^>]+http-equiv=["\']?Content-Type["\']?[^>]+content=["\']?[^"\']*charset=([^"\'>\s;]+)',
            head_text, re.I
        )
        if content_type_match:
            charset = content_type_match.group(1)
            try:
                return content_bytes.decode(charset)
            except (UnicodeDecodeError, LookupError) as e:
                print(f"DEBUG: Meta Content-Type charset {charset} failed: {e}", file=sys.stderr)
    except Exception:
        pass

    # 3. Try charset_normalizer for smart detection
    try:
        from charset_normalizer import from_bytes
        result = from_bytes(content_bytes).best()
        if result:
            print(f"DEBUG: charset_normalizer detected: {result.encoding}", file=sys.stderr)
            return str(result)
    except ImportError:
        pass
    except Exception as e:
        print(f"DEBUG: charset_normalizer failed: {e}", file=sys.stderr)

    # 4. Fall back to UTF-8 with replacement (never ignore errors!)
    # Using 'replace' instead of 'ignore' to make encoding issues visible
    print(f"DEBUG: Falling back to UTF-8 decode", file=sys.stderr)
    return content_bytes.decode('utf-8', errors='replace')


def _trafilatura_extraction_ladder(html: str, soup: BeautifulSoup) -> tuple[list[str], str]:
    """
    Multi-stage trafilatura extraction with automatic fallbacks.

    EXTRACTION LADDER (in order of precision -> recall):
    1. trafilatura.extract() with favor_recall=True (balanced)
    2. trafilatura.extract() with favor_precision=False (more recall)
    3. trafilatura.bare_extraction() baseline (maximum recall)
    4. BeautifulSoup fallback (DOM-based, captures everything)

    Each stage checks extraction quality. If insufficient, escalates to next.

    Args:
        html: Raw HTML content.
        soup: BeautifulSoup parsed HTML (for fallback).

    Returns:
        Tuple of (content_blocks, extraction_method_used).
    """
    import sys

    # Stage 1: Balanced extraction with favor_recall
    try:
        content = trafilatura.extract(
            html,
            include_comments=False,
            include_tables=True,
            include_links=False,
            include_images=False,
            favor_recall=True,
            favor_precision=False,
        )
        if content:
            blocks = _split_to_blocks(content)
            if _extraction_quality_ok(blocks, "trafilatura_favor_recall"):
                return blocks, "trafilatura_favor_recall"
            print(f"DEBUG: Stage 1 insufficient ({len(blocks)} blocks), escalating...", file=sys.stderr)
    except Exception as e:
        print(f"DEBUG: Stage 1 failed: {e}", file=sys.stderr)

    # Stage 2: Maximum recall settings
    try:
        content = trafilatura.extract(
            html,
            include_comments=False,
            include_tables=True,
            include_links=True,  # Include more
            include_images=False,
            favor_recall=True,
            favor_precision=False,
            no_fallback=False,
        )
        if content:
            blocks = _split_to_blocks(content)
            if _extraction_quality_ok(blocks, "trafilatura_max_recall"):
                return blocks, "trafilatura_max_recall"
            print(f"DEBUG: Stage 2 insufficient ({len(blocks)} blocks), escalating...", file=sys.stderr)
    except Exception as e:
        print(f"DEBUG: Stage 2 failed: {e}", file=sys.stderr)

    # Stage 3: Try bare_extraction for raw content (if available)
    try:
        if hasattr(trafilatura, 'bare_extraction'):
            result = trafilatura.bare_extraction(html)
            if result and isinstance(result, dict):
                # bare_extraction returns a dict with 'text' key
                content = result.get('text', '')
                if content:
                    blocks = _split_to_blocks(content)
                    if _extraction_quality_ok(blocks, "trafilatura_bare"):
                        return blocks, "trafilatura_bare"
                    print(f"DEBUG: Stage 3 insufficient ({len(blocks)} blocks), escalating...", file=sys.stderr)
    except Exception as e:
        print(f"DEBUG: Stage 3 failed: {e}", file=sys.stderr)

    # Stage 4: Try html2txt for maximum text extraction (if available)
    try:
        if hasattr(trafilatura, 'html2txt'):
            content = trafilatura.html2txt(html)
            if content:
                blocks = _split_to_blocks(content)
                if _extraction_quality_ok(blocks, "trafilatura_html2txt"):
                    return blocks, "trafilatura_html2txt"
                print(f"DEBUG: Stage 4 insufficient ({len(blocks)} blocks), escalating...", file=sys.stderr)
    except Exception as e:
        print(f"DEBUG: Stage 4 failed: {e}", file=sys.stderr)

    # Stage 5: Final fallback - BeautifulSoup DOM extraction
    print(f"DEBUG: All trafilatura methods insufficient, using BeautifulSoup fallback", file=sys.stderr)
    blocks = _extract_content_fallback(soup)
    return blocks, "beautifulsoup_fallback"


def _split_to_blocks(content: str) -> list[str]:
    """Split extracted content into paragraph blocks."""
    if not content:
        return []
    # Split on double newlines (paragraph boundaries)
    paragraphs = [p.strip() for p in content.split("\n\n") if p.strip()]
    # Also handle single newlines for list items etc.
    if len(paragraphs) < 3:
        paragraphs = [p.strip() for p in content.split("\n") if p.strip() and len(p.strip()) > 10]
    return paragraphs


def _extraction_quality_ok(blocks: list[str], method: str) -> bool:
    """
    Check if extraction quality meets minimum thresholds.

    Args:
        blocks: Extracted content blocks.
        method: Extraction method name (for logging).

    Returns:
        True if quality is acceptable.
    """
    if not blocks:
        return False

    # Count total words
    total_words = sum(len(block.split()) for block in blocks)

    # Count substantial blocks (more than 5 words)
    substantial_blocks = sum(1 for block in blocks if len(block.split()) > 5)

    # Check thresholds
    if total_words < MIN_WORD_COUNT:
        return False

    if substantial_blocks < MIN_BLOCK_COUNT:
        return False

    return True


def fetch_url_content(
    url: str,
    timeout: int = 30,
    use_firecrawl: bool = True,
    firecrawl_api_key: Optional[str] = None,
) -> PageMeta:
    """
    Fetch and extract content from a URL.

    Uses FireCrawl for structured content extraction when available,
    falling back to trafilatura/BeautifulSoup otherwise.

    Args:
        url: The URL to fetch content from.
        timeout: Request timeout in seconds.
        use_firecrawl: Whether to attempt FireCrawl extraction first.
        firecrawl_api_key: Optional FireCrawl API key.

    Returns:
        PageMeta object containing extracted content and metadata.

    Raises:
        ContentExtractionError: If fetching or parsing fails.
    """
    # Validate URL
    parsed = urlparse(url)
    if not parsed.scheme or not parsed.netloc:
        raise ContentExtractionError(f"Invalid URL: {url}")

    # Try FireCrawl first if enabled and API key is available
    if use_firecrawl:
        api_key = firecrawl_api_key or os.environ.get("FIRECRAWL_API_KEY")
        if api_key:
            try:
                from .firecrawl_client import FireCrawlClient, FireCrawlError

                client = FireCrawlClient(api_key=api_key)
                if client.is_available:
                    return client.scrape_url(url, timeout=timeout)
            except (ImportError, FireCrawlError) as e:
                # Fall back to traditional extraction
                pass

    # Fall back to traditional extraction

    try:
        response = requests.get(url, headers=DEFAULT_HEADERS, timeout=timeout)
        response.raise_for_status()
    except requests.RequestException as e:
        raise ContentExtractionError(f"Failed to fetch URL: {e}")

    # CRITICAL: Proper encoding detection to prevent mojibake
    # Do NOT rely on response.text which can use wrong encoding
    html = _decode_html_safely(response)

    # Parse HTML with BeautifulSoup for meta extraction
    soup = BeautifulSoup(html, "lxml")

    # Extract title - use separator to preserve spaces between nested elements
    title = None
    title_tag = soup.find("title")
    if title_tag:
        title = title_tag.get_text(separator=" ", strip=True)
        title = " ".join(title.split())  # Normalize whitespace

    # Extract meta description
    meta_description = None
    meta_desc_tag = soup.find("meta", attrs={"name": "description"})
    if meta_desc_tag and meta_desc_tag.get("content"):
        meta_description = meta_desc_tag["content"].strip()
        meta_description = " ".join(meta_description.split())  # Normalize whitespace

    # Extract H1 - use separator to preserve spaces between nested elements
    # This prevents "Park Ave" + "premium" becoming "Park Avepremium"
    h1 = None
    h1_tag = soup.find("h1")
    if h1_tag:
        h1 = h1_tag.get_text(separator=" ", strip=True)
        h1 = " ".join(h1.split())  # Normalize whitespace

    # Extract main content using EXTRACTION LADDER approach
    # This ensures maximum content recall with multiple fallback strategies
    content_blocks: list[str] = []
    extraction_method = "none"

    if trafilatura:
        content_blocks, extraction_method = _trafilatura_extraction_ladder(html, soup)
    else:
        # No trafilatura - use BeautifulSoup fallback
        content_blocks = _extract_content_fallback(soup)
        extraction_method = "beautifulsoup_fallback"

    import sys
    print(f"DEBUG: Extraction method used: {extraction_method}, blocks: {len(content_blocks)}", file=sys.stderr)

    # CRITICAL: Apply text repair to fix encoding issues and mojibake
    # This prevents corruption like "â€™" appearing instead of apostrophes
    title_repaired, _ = repair_text(title) if title else (title, False)
    meta_description_repaired, _ = repair_text(meta_description) if meta_description else (meta_description, False)
    h1_repaired, _ = repair_text(h1) if h1 else (h1, False)
    content_blocks_repaired, corruption_count = repair_content_blocks(content_blocks)

    if corruption_count > 0:
        import sys
        print(f"WARNING: Repaired encoding corruption in {corruption_count} content blocks", file=sys.stderr)

    return PageMeta(
        title=title_repaired,
        meta_description=meta_description_repaired,
        h1=h1_repaired,
        content_blocks=content_blocks_repaired,
        url=url,
    )


def _extract_content_fallback(soup: BeautifulSoup) -> list[str]:
    """
    Fallback content extraction when trafilatura is not available.

    ENHANCED FOR HOMEPAGE/LANDING PAGES:
    - Does NOT drop short blocks (homepages are made of short blocks)
    - Extracts all text nodes including CTAs, stats, testimonials
    - Preserves cards, features, and promotional content
    - Includes address/contact/legal blocks

    Args:
        soup: BeautifulSoup parsed HTML.

    Returns:
        List of content blocks (paragraphs).
    """
    content_blocks: list[str] = []

    # Remove script, style, and truly navigational elements
    # NOTE: Keep header/footer for homepages as they may contain key content
    for tag in soup.find_all(["script", "style", "noscript", "iframe"]):
        tag.decompose()

    # Look for main content container but fall back to body for landing pages
    main_content = (
        soup.find("main")
        or soup.find("article")
        or soup.find("div", class_=re.compile(r"content|post|entry|article|body|wrapper|container", re.I))
        or soup.find("div", id=re.compile(r"content|post|entry|article|main|app", re.I))
    )

    if main_content:
        source = main_content
    else:
        source = soup.body or soup

    # Extract ALL text-bearing elements (not just p/h/li)
    # This captures cards, stats, CTAs, testimonials common on homepages
    text_elements = [
        "p", "h1", "h2", "h3", "h4", "h5", "h6", "li",
        "span", "div", "a", "button", "label",  # Interactive/card elements
        "blockquote", "figcaption", "cite",  # Quotes/testimonials
        "address", "time",  # Contact/date info
        "td", "th", "caption",  # Table content
    ]

    seen_texts = set()  # Deduplicate exact matches

    for element in source.find_all(text_elements):
        # Get direct text content (avoid nested duplication)
        text = element.get_text(separator=" ", strip=True)

        if not text:
            continue

        # Normalize whitespace
        text = " ".join(text.split())

        # Skip if already seen (exact match)
        if text in seen_texts:
            continue

        # IMPORTANT: Do NOT skip short blocks for homepages
        # Short blocks include: CTAs, stats, feature headlines, nav items
        # Only skip truly trivial content (< 3 chars or just punctuation)
        if len(text) < 3 or text.strip() in [".", ",", "-", "•", "|"]:
            continue

        # Skip pure navigation/social links
        if element.name == "a" and element.get("href", "").startswith(("#", "javascript:", "mailto:", "tel:")):
            # But keep if it has substantial text (could be a CTA)
            if len(text) < 10:
                continue

        seen_texts.add(text)
        content_blocks.append(text)

    return content_blocks


def load_docx_content(file_path: Union[str, Path]) -> DocxContent:
    """
    Load and extract content from a Word document.

    Args:
        file_path: Path to the .docx file.

    Returns:
        DocxContent object containing extracted paragraphs and headings.

    Raises:
        ContentExtractionError: If the file cannot be read or parsed.
    """
    path = Path(file_path)

    if not path.exists():
        raise ContentExtractionError(f"File not found: {file_path}")

    if not path.suffix.lower() == ".docx":
        raise ContentExtractionError(f"File must be a .docx file: {file_path}")

    try:
        doc = Document(str(path))
    except Exception as e:
        raise ContentExtractionError(f"Failed to open Word document: {e}")

    paragraphs: list[ParagraphBlock] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Determine heading level from style
        heading_level = HeadingLevel.BODY
        style_name = None

        if para.style:
            style_name = para.style.name
            if style_name in HEADING_STYLE_MAP:
                heading_level = HEADING_STYLE_MAP[style_name]
            elif style_name and style_name.startswith("Heading"):
                # Try to parse numeric heading level
                try:
                    level = int(style_name.replace("Heading ", "").strip())
                    if 1 <= level <= 6:
                        heading_level = HeadingLevel(level)
                except ValueError:
                    pass

        paragraphs.append(
            ParagraphBlock(
                text=text,
                heading_level=heading_level,
                style_name=style_name,
            )
        )

    return DocxContent(paragraphs=paragraphs, source_path=str(path))


def load_content(source: str) -> Union[PageMeta, DocxContent]:
    """
    Load content from either a URL or a file path.

    Args:
        source: URL or file path to load content from.

    Returns:
        PageMeta (for URLs) or DocxContent (for .docx files).

    Raises:
        ContentExtractionError: If the source is invalid or cannot be loaded.
    """
    # Check if it's a URL
    parsed = urlparse(source)
    if parsed.scheme in ("http", "https"):
        return fetch_url_content(source)

    # Check if it's a file path
    path = Path(source)
    if path.suffix.lower() == ".docx":
        return load_docx_content(path)

    raise ContentExtractionError(
        f"Invalid source: {source}. Must be a URL (http/https) or a .docx file path."
    )


def convert_page_meta_to_blocks(page_meta: PageMeta) -> list[ParagraphBlock]:
    """
    Convert PageMeta content blocks to ParagraphBlock list for uniform processing.

    Args:
        page_meta: PageMeta object from URL extraction.

    Returns:
        List of ParagraphBlock objects.
    """
    blocks: list[ParagraphBlock] = []

    # Add H1 if present
    h1_text = page_meta.h1.strip() if page_meta.h1 else ""
    if h1_text:
        blocks.append(ParagraphBlock(text=h1_text, heading_level=HeadingLevel.H1))

    # Add content blocks as body paragraphs
    # In real scenarios, we might want to detect headings within content_blocks
    for block in page_meta.content_blocks:
        block_text = block.strip()

        # Skip empty blocks
        if not block_text:
            continue

        # Skip blocks that match the H1 text (avoid duplication from FireCrawl)
        if h1_text and block_text == h1_text:
            continue

        # Simple heuristic: short text that looks like a heading
        if len(block_text) < 100 and not block_text.endswith((".", "!", "?")):
            # Might be a heading - default to H2
            blocks.append(ParagraphBlock(text=block_text, heading_level=HeadingLevel.H2))
        else:
            blocks.append(ParagraphBlock(text=block_text, heading_level=HeadingLevel.BODY))

    return blocks


# =============================================================================
# V2 ARCHITECTURE: ContentDocument Extraction
# =============================================================================
# These functions produce ContentDocument with typed blocks, preserving
# structure for tables, lists, headings, and paragraphs.
# =============================================================================


def _html_tag_to_block_type(tag_name: str) -> BlockType:
    """Map HTML tag name to BlockType."""
    tag_map: dict[str, BlockType] = {
        "h1": "h1",
        "h2": "h2",
        "h3": "h3",
        "h4": "h4",
        "h5": "h5",
        "h6": "h6",
        "p": "p",
        "ul": "ul",
        "ol": "ol",
        "li": "li",
        "table": "table",
        "tr": "tr",
        "td": "td",
        "th": "th",
        "caption": "caption",
        "blockquote": "blockquote",
        "img": "image",
        "hr": "hr",
    }
    return tag_map.get(tag_name.lower(), "p")


def _parse_html_to_blocks(soup: BeautifulSoup, source_element=None) -> list[Block]:
    """
    Parse HTML content into typed Block objects.

    Preserves structure for headings, paragraphs, lists, and tables.

    Args:
        soup: BeautifulSoup parsed HTML.
        source_element: Optional element to extract from (defaults to main content).

    Returns:
        List of Block objects representing the document structure.
    """
    blocks: list[Block] = []

    # Remove unwanted elements
    for tag in soup.find_all(["script", "style", "nav", "footer", "aside", "noscript"]):
        tag.decompose()

    # Find main content container
    if source_element is None:
        source_element = (
            soup.find("main")
            or soup.find("article")
            or soup.find("div", class_=re.compile(r"content|post|entry|article|body", re.I))
            or soup.find("div", id=re.compile(r"content|post|entry|article|main", re.I))
            or soup.body
            or soup
        )

    # Process elements in order
    block_tags = ["h1", "h2", "h3", "h4", "h5", "h6", "p", "ul", "ol", "table", "blockquote", "hr"]

    for element in source_element.find_all(block_tags, recursive=True):
        # Skip elements that are nested inside other tracked elements (handled by parent)
        parent = element.parent
        skip = False
        while parent and parent != source_element:
            if parent.name in ["ul", "ol", "table", "blockquote"]:
                skip = True
                break
            parent = parent.parent
        if skip:
            continue

        block = _element_to_block(element)
        if block:
            blocks.append(block)

    return blocks


def _element_to_block(element) -> Optional[Block]:
    """
    Convert a BeautifulSoup element to a Block.

    Args:
        element: BeautifulSoup element.

    Returns:
        Block object or None if element should be skipped.
    """
    tag_name = element.name.lower()

    # Handle headings and paragraphs
    if tag_name in ("h1", "h2", "h3", "h4", "h5", "h6", "p"):
        text = element.get_text(strip=True)
        if not text or len(text) < 3:
            return None
        block_type = _html_tag_to_block_type(tag_name)
        return Block.create(block_type=block_type, text=text)

    # Handle unordered/ordered lists
    if tag_name in ("ul", "ol"):
        children: list[Block] = []
        for li in element.find_all("li", recursive=False):
            li_text = li.get_text(strip=True)
            if li_text:
                li_block = Block.create(block_type="li", text=li_text)
                children.append(li_block)

        if not children:
            return None

        block_type: BlockType = "ul" if tag_name == "ul" else "ol"
        return Block.create(
            block_type=block_type,
            children=children,
            attrs={"item_count": len(children)},
        )

    # Handle tables
    if tag_name == "table":
        rows: list[Block] = []
        caption_text = None

        # Extract caption if present
        caption = element.find("caption")
        if caption:
            caption_text = caption.get_text(strip=True)

        # Extract rows
        for tr in element.find_all("tr"):
            cells: list[Block] = []
            for cell in tr.find_all(["td", "th"]):
                cell_type: BlockType = "th" if cell.name == "th" else "td"
                cell_text = cell.get_text(strip=True)
                cell_block = Block.create(block_type=cell_type, text=cell_text)
                cells.append(cell_block)

            if cells:
                row_block = Block.create(block_type="tr", children=cells)
                rows.append(row_block)

        if not rows:
            return None

        attrs: dict = {"row_count": len(rows)}
        if caption_text:
            attrs["caption"] = caption_text

        return Block.create(
            block_type="table",
            children=rows,
            attrs=attrs,
        )

    # Handle blockquotes
    if tag_name == "blockquote":
        text = element.get_text(strip=True)
        if not text:
            return None
        return Block.create(block_type="blockquote", text=text)

    # Handle horizontal rules
    if tag_name == "hr":
        return Block.create(block_type="hr")

    return None


def fetch_url_as_document(
    url: str,
    timeout: int = 30,
    use_firecrawl: bool = True,
    firecrawl_api_key: Optional[str] = None,
) -> ContentDocument:
    """
    Fetch and extract content from a URL into a ContentDocument.

    V2 Architecture: Returns typed blocks preserving document structure
    including tables, lists, and proper heading hierarchy.

    Args:
        url: The URL to fetch content from.
        timeout: Request timeout in seconds.
        use_firecrawl: Whether to attempt FireCrawl extraction first.
        firecrawl_api_key: Optional FireCrawl API key.

    Returns:
        ContentDocument object with typed blocks.

    Raises:
        ContentExtractionError: If fetching or parsing fails.
    """
    # Validate URL
    parsed = urlparse(url)
    if not parsed.scheme or not parsed.netloc:
        raise ContentExtractionError(f"Invalid URL: {url}")

    # Try FireCrawl first if enabled
    if use_firecrawl:
        api_key = firecrawl_api_key or os.environ.get("FIRECRAWL_API_KEY")
        if api_key:
            try:
                from .firecrawl_client import FireCrawlClient, FireCrawlError

                client = FireCrawlClient(api_key=api_key)
                if client.is_available:
                    page_meta = client.scrape_url(url, timeout=timeout)
                    # Convert PageMeta to ContentDocument
                    return _page_meta_to_document(page_meta, url)
            except (ImportError, FireCrawlError):
                pass

    # Traditional extraction with structure preservation
    try:
        response = requests.get(url, headers=DEFAULT_HEADERS, timeout=timeout)
        response.raise_for_status()
    except requests.RequestException as e:
        raise ContentExtractionError(f"Failed to fetch URL: {e}")

    html = response.text
    soup = BeautifulSoup(html, "lxml")

    # Extract meta information
    title = None
    title_tag = soup.find("title")
    if title_tag:
        title = title_tag.get_text(strip=True)

    meta_description = None
    meta_desc_tag = soup.find("meta", attrs={"name": "description"})
    if meta_desc_tag and meta_desc_tag.get("content"):
        meta_description = meta_desc_tag["content"].strip()

    # Parse content into typed blocks
    blocks = _parse_html_to_blocks(soup)

    # Add meta blocks at the start
    meta_blocks: list[Block] = []
    if title:
        meta_blocks.append(Block.create(block_type="meta_title", text=title))
    if meta_description:
        meta_blocks.append(Block.create(block_type="meta_desc", text=meta_description))

    all_blocks = meta_blocks + blocks

    return ContentDocument(
        blocks=all_blocks,
        source_url=url,
        extracted_title=title,
        extracted_meta_desc=meta_description,
        extraction_timestamp=datetime.now().isoformat(),
    )


def _page_meta_to_document(page_meta: PageMeta, url: str) -> ContentDocument:
    """
    Convert a PageMeta object to ContentDocument.

    Used when FireCrawl extraction returns PageMeta format.

    Args:
        page_meta: PageMeta from FireCrawl or other extraction.
        url: Source URL.

    Returns:
        ContentDocument with typed blocks.
    """
    blocks: list[Block] = []

    # Add meta blocks
    if page_meta.title:
        blocks.append(Block.create(block_type="meta_title", text=page_meta.title))
    if page_meta.meta_description:
        blocks.append(Block.create(block_type="meta_desc", text=page_meta.meta_description))

    # Add H1 block
    h1_text = page_meta.h1.strip() if page_meta.h1 else ""
    if h1_text:
        blocks.append(Block.create(block_type="h1", text=h1_text))

    # Convert content blocks to typed blocks with heuristic detection
    for content in page_meta.content_blocks:
        content = content.strip()
        if not content:
            continue

        # Skip blocks that match the H1 text (avoid duplication from FireCrawl)
        if h1_text and content == h1_text:
            continue

        # Detect block type from content characteristics
        block_type = _infer_block_type_from_text(content)
        blocks.append(Block.create(block_type=block_type, text=content))

    return ContentDocument(
        blocks=blocks,
        source_url=url,
        extracted_title=page_meta.title,
        extracted_meta_desc=page_meta.meta_description,
        extraction_timestamp=datetime.now().isoformat(),
    )


def _infer_block_type_from_text(text: str) -> BlockType:
    """
    Infer the block type from text content using heuristics.

    Args:
        text: Content text.

    Returns:
        Inferred BlockType.
    """
    # Short text without terminal punctuation likely a heading
    if len(text) < 100 and not text.endswith((".", "!", "?", ":")):
        # Very short = likely H2, slightly longer = H3
        if len(text) < 50:
            return "h2"
        return "h3"

    # Text starting with bullet indicators
    if text.startswith(("• ", "- ", "* ", "· ")):
        return "li"

    # Text starting with numbers and period/parenthesis (ordered list item)
    if re.match(r"^\d+[\.\)]\s", text):
        return "li"

    # Default to paragraph
    return "p"


def load_docx_as_document(file_path: Union[str, Path]) -> ContentDocument:
    """
    Load and extract content from a Word document into ContentDocument.

    V2 Architecture: Preserves document structure including:
    - Heading hierarchy (H1-H6)
    - Paragraphs with run-level formatting
    - Tables with cell structure
    - Lists (bullet and numbered)

    Args:
        file_path: Path to the .docx file.

    Returns:
        ContentDocument with typed blocks and preserved formatting.

    Raises:
        ContentExtractionError: If the file cannot be read or parsed.
    """
    path = Path(file_path)

    if not path.exists():
        raise ContentExtractionError(f"File not found: {file_path}")

    if not path.suffix.lower() == ".docx":
        raise ContentExtractionError(f"File must be a .docx file: {file_path}")

    try:
        doc = Document(str(path))
    except Exception as e:
        raise ContentExtractionError(f"Failed to open Word document: {e}")

    blocks: list[Block] = []
    extracted_title: Optional[str] = None

    # Process paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Determine block type from style
        block_type = _docx_style_to_block_type(para.style.name if para.style else None)

        # Capture title (first H1)
        if block_type == "h1" and extracted_title is None:
            extracted_title = text

        # Extract runs for formatting preservation
        runs = _extract_docx_runs(para)

        block = Block.create(
            block_type=block_type,
            text=text,
            runs=runs if runs else None,
            attrs={"style_name": para.style.name if para.style else None},
        )
        blocks.append(block)

    # Process tables
    for table in doc.tables:
        table_block = _extract_docx_table(table)
        if table_block:
            blocks.append(table_block)

    return ContentDocument(
        blocks=blocks,
        source_docx_path=str(path),
        extracted_title=extracted_title,
        extraction_timestamp=datetime.now().isoformat(),
    )


def _docx_style_to_block_type(style_name: Optional[str]) -> BlockType:
    """
    Map Word document style name to BlockType.

    Args:
        style_name: Word style name.

    Returns:
        Corresponding BlockType.
    """
    if not style_name:
        return "p"

    style_lower = style_name.lower()

    # Handle heading styles
    if style_lower.startswith("heading"):
        try:
            level = int(style_lower.replace("heading", "").strip())
            if 1 <= level <= 6:
                return f"h{level}"  # type: ignore
        except ValueError:
            pass
        return "h2"  # Default for unrecognized heading

    if style_lower == "title":
        return "h1"

    if "list" in style_lower or "bullet" in style_lower:
        return "li"

    if "quote" in style_lower:
        return "blockquote"

    return "p"


def _extract_docx_runs(paragraph) -> list[Run]:
    """
    Extract formatted runs from a DOCX paragraph.

    Args:
        paragraph: python-docx paragraph object.

    Returns:
        List of Run objects with formatting preserved.
    """
    runs: list[Run] = []

    for run in paragraph.runs:
        if not run.text:
            continue

        runs.append(Run(
            text=run.text,
            bold=run.bold or False,
            italic=run.italic or False,
            underline=run.underline is not None and run.underline,
            strike=run.font.strike or False,
            font_name=run.font.name,
            font_size=run.font.size.pt if run.font.size else None,
        ))

    return runs


def _extract_docx_table(table) -> Optional[Block]:
    """
    Extract a table from DOCX into a Block structure.

    Args:
        table: python-docx table object.

    Returns:
        Block representing the table, or None if empty.
    """
    rows: list[Block] = []

    for row in table.rows:
        cells: list[Block] = []
        is_header_row = False

        for idx, cell in enumerate(row.cells):
            cell_text = cell.text.strip()

            # Detect header row (first row or cells with bold text)
            if idx == 0 and cell.paragraphs:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.bold:
                            is_header_row = True
                            break

            cell_type: BlockType = "th" if is_header_row else "td"
            cell_block = Block.create(block_type=cell_type, text=cell_text)
            cells.append(cell_block)

        if cells:
            row_block = Block.create(block_type="tr", children=cells)
            rows.append(row_block)

    if not rows:
        return None

    return Block.create(
        block_type="table",
        children=rows,
        attrs={"row_count": len(rows)},
    )


def load_content_as_document(source: str) -> ContentDocument:
    """
    Load content from either a URL or a file path into ContentDocument.

    V2 Architecture: Unified loader that returns ContentDocument with
    typed blocks regardless of source type.

    Args:
        source: URL or file path to load content from.

    Returns:
        ContentDocument with typed blocks.

    Raises:
        ContentExtractionError: If the source is invalid or cannot be loaded.
    """
    # Check if it's a URL
    parsed = urlparse(source)
    if parsed.scheme in ("http", "https"):
        return fetch_url_as_document(source)

    # Check if it's a file path
    path = Path(source)
    if path.suffix.lower() == ".docx":
        return load_docx_as_document(path)

    raise ContentExtractionError(
        f"Invalid source: {source}. Must be a URL (http/https) or a .docx file path."
    )
