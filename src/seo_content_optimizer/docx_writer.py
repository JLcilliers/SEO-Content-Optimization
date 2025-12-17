"""
Word document writer with green highlight support.

This module generates the final optimized .docx output with:
- Reading guide at the top
- Current vs Optimized Meta Elements table
- Optimized content with green-highlighted changes
- Poppins font throughout with consistent spacing
"""

import re
from pathlib import Path
from typing import Optional, Union

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, Twips
from docx.text.paragraph import Paragraph

from .llm_client import ADD_END, ADD_START
from .diff_markers import normalize_paragraph_spacing
from .models import (
    ContentBlock,
    ContentBlockType,
    FAQItem,
    HeadingLevel,
    ListBlock,
    MetaElement,
    OptimizationResult,
    ParagraphBlock,
    TableBlock,
)


# Marker constants for identifying changed content
MARK_START = ADD_START  # "[[[ADD]]]"
MARK_END = ADD_END  # "[[[ENDADD]]]"

# Regex pattern for invalid XML 1.0 characters
# Valid XML 1.0 chars: #x9 | #xA | #xD | [#x20-#xD7FF] | [#xE000-#xFFFD] | [#x10000-#x10FFFF]
# We remove everything else (control characters 0x00-0x08, 0x0B, 0x0C, 0x0E-0x1F)
_INVALID_XML_CHARS_RE = re.compile(
    r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]"
)


def sanitize_for_xml(text: str) -> str:
    """
    Remove invalid XML characters from text to prevent Word "unreadable content" errors.

    XML 1.0 does not allow certain control characters. This function removes them
    to ensure the generated DOCX is valid and opens without repair dialogs.

    Args:
        text: Input text that may contain invalid XML characters.

    Returns:
        Sanitized text safe for XML/DOCX.
    """
    if not text:
        return text
    return _INVALID_XML_CHARS_RE.sub("", text)


# MARKER LEAKAGE PREVENTION
# These patterns detect markers that may have leaked into output
MARKER_LEAK_PATTERNS = [
    r"\[\[\[ADD\]\]\]",      # Main marker start
    r"\[\[\[ENDADD\]\]\]",   # Main marker end
    r"add\.\.\.endadd",       # Corrupted marker pattern
    r"\[add\]",              # Lowercase variant
    r"\[/add\]",             # HTML-style variant
    r"<<<ADD>>>",            # Alternative marker style
    r"<<<ENDADD>>>",         # Alternative marker style
]


def validate_no_marker_leakage(text: str) -> tuple[bool, list[str]]:
    """
    Validate that text doesn't contain leaked marker tokens.

    CRITICAL: This is a final safety check to ensure no markers
    appear in visible output text.

    Args:
        text: Text to validate.

    Returns:
        Tuple of (is_valid, list of found leak patterns).
    """
    if not text:
        return True, []

    found_leaks = []
    text_lower = text.lower()

    for pattern in MARKER_LEAK_PATTERNS:
        if re.search(pattern, text, re.IGNORECASE):
            found_leaks.append(pattern)

    # Also check for partial markers
    partial_markers = [
        "[[[", "]]]",  # Bracket patterns
        "ENDADD", "[[ADD",  # Partial marker text
    ]
    for marker in partial_markers:
        if marker in text:
            found_leaks.append(f"partial marker: {marker}")

    return len(found_leaks) == 0, found_leaks


def strip_leaked_markers(text: str) -> str:
    """
    Remove any leaked markers from text as a safety net.

    CRITICAL: This function must NOT strip boundary whitespace!
    Stripping causes word concatenation bugs like "Avepremium"
    when markers are between words.

    This should only catch edge cases where markers weren't
    properly processed. The main processing should handle markers
    correctly, but this provides a backup.

    Args:
        text: Text that may contain leaked markers.

    Returns:
        Cleaned text with markers removed (whitespace preserved).
    """
    if not text:
        return text

    result = text

    # Remove known marker patterns (preserve surrounding whitespace!)
    result = result.replace(MARK_START, "")
    result = result.replace(MARK_END, "")

    # Remove any partial/corrupted marker patterns
    for pattern in MARKER_LEAK_PATTERNS:
        result = re.sub(pattern, "", result, flags=re.IGNORECASE)

    # Clean up resulting double/triple spaces -> single space
    # But do NOT strip leading/trailing - that causes boundary bugs!
    result = re.sub(r"  +", " ", result)

    # IMPORTANT: Do NOT call .strip() here!
    # Stripping removes boundary whitespace and causes word concatenation
    # like "in green" becoming "ingreen" or "Ave premium" -> "Avepremium"
    return result


def set_cell_shading(cell, color: str) -> None:
    """Set background color/shading for a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tcPr.append(shd)


def _normalize_markers_for_newlines(text: str) -> str:
    """
    Ensure markers don't span newlines by splitting them at line boundaries.

    If a marker contains a newline, split it so each line has its own markers.
    This prevents markers from being broken when text is split on newlines.

    Example:
        "[[[ADD]]]line1\nline2[[[ENDADD]]]"
        becomes:
        "[[[ADD]]]line1[[[ENDADD]]]\n[[[ADD]]]line2[[[ENDADD]]]"

    Args:
        text: Text with markers that may span newlines.

    Returns:
        Text with markers normalized so they don't span newlines.
    """
    if MARK_START not in text:
        return text

    result = []
    pos = 0

    while pos < len(text):
        start = text.find(MARK_START, pos)
        if start == -1:
            # No more markers - add rest of text
            result.append(text[pos:])
            break

        # Add text before the marker
        result.append(text[pos:start])

        # Find closing marker
        end = text.find(MARK_END, start)
        if end == -1:
            # Malformed: no closing marker - add rest of text
            result.append(text[start:])
            break

        # Extract content between markers
        content_start = start + len(MARK_START)
        content = text[content_start:end]

        # Check if content contains newlines
        if "\n" in content:
            # Split content on newlines and wrap each part
            lines = content.split("\n")
            marked_lines = []
            for i, line in enumerate(lines):
                if line:  # Only wrap non-empty lines
                    marked_lines.append(f"{MARK_START}{line}{MARK_END}")
                if i < len(lines) - 1:
                    # Add newline between parts (outside markers)
                    marked_lines.append("\n")
            result.append("".join(marked_lines))
        else:
            # No newlines - keep marker as-is
            result.append(f"{MARK_START}{content}{MARK_END}")

        pos = end + len(MARK_END)

    return "".join(result)


def add_marked_text(paragraph: Paragraph, text: str, font_name: str = "Poppins") -> None:
    """
    Write `text` into `paragraph`, converting [[[ADD]]]/[[[ENDADD]]] segments
    into BRIGHT_GREEN highlighted runs, and leaving all other text normal.

    This is the single canonical way to add optimized text with highlights.
    Text is sanitized to remove invalid XML characters before insertion.

    MARKER LEAKAGE PREVENTION: Malformed markers are stripped, not outputted.

    Args:
        paragraph: The paragraph to add text to.
        text: Text with optional [[[ADD]]]...[[[ENDADD]]] markers.
        font_name: Font to use for all runs (default: Poppins).
    """
    # Sanitize text to remove invalid XML characters that cause Word errors
    text = sanitize_for_xml(text)

    # Normalize markers so they don't span newlines
    # (This helps when text is later split on newlines for structure)
    text = _normalize_markers_for_newlines(text)

    while text:
        start = text.find(MARK_START)
        if start == -1:
            # No more markers - add remaining text as plain
            # SAFETY: Validate no leaked markers in plain text
            if text:
                clean_text = strip_leaked_markers(text)
                if clean_text:
                    run = paragraph.add_run(clean_text)
                    run.font.name = font_name
            break

        # Add plain prefix before the marker
        if start > 0:
            # SAFETY: Strip any leaked markers from plain text
            prefix = strip_leaked_markers(text[:start])
            if prefix:
                run = paragraph.add_run(prefix)
                run.font.name = font_name

        # Find closing marker
        end = text.find(MARK_END, start)
        if end == -1:
            # Malformed: no closing marker
            # SAFETY: Strip the marker text and output as plain
            remaining = text[start + len(MARK_START):]  # Skip the malformed start marker
            clean_remaining = strip_leaked_markers(remaining)
            if clean_remaining:
                run = paragraph.add_run(clean_remaining)
                run.font.name = font_name
            break

        # Extract and add highlighted content
        highlighted = text[start + len(MARK_START): end]
        if highlighted:
            # SAFETY: Validate highlighted content has no nested markers
            clean_highlighted = strip_leaked_markers(highlighted)
            if clean_highlighted:
                run = paragraph.add_run(clean_highlighted)
                run.font.name = font_name
                # CRITICAL: Use the official enum for proper Word compatibility
                run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN

        # Continue after the closing marker
        text = text[end + len(MARK_END):]


class DocxWriter:
    """
    Writes optimization results to a Word document.

    Creates a professional document with:
    - Reading guide explaining the green highlight convention
    - Table comparing current vs optimized meta elements
    - Full optimized content with changes highlighted
    """

    def __init__(self):
        """Initialize the document writer."""
        self.doc = Document()
        self._setup_styles()

    def _setup_styles(self) -> None:
        """Configure document styles with Poppins font and consistent spacing."""
        # Font settings
        FONT_NAME = "Poppins"
        BODY_SIZE = Pt(11)

        # Spacing settings (in points)
        PARA_SPACE_BEFORE = Pt(6)
        PARA_SPACE_AFTER = Pt(6)
        LINE_SPACING = 1.15  # 115% line spacing for readability

        # Configure Normal style (base for all text)
        normal_style = self.doc.styles["Normal"]
        normal_style.font.name = FONT_NAME
        normal_style.font.size = BODY_SIZE
        normal_style.paragraph_format.space_before = PARA_SPACE_BEFORE
        normal_style.paragraph_format.space_after = PARA_SPACE_AFTER
        normal_style.paragraph_format.line_spacing = LINE_SPACING

        # Set font for East Asian text fallback (required for proper font embedding)
        normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)

        # Configure Title style
        if "Title" in self.doc.styles:
            title_style = self.doc.styles["Title"]
            title_style.font.name = FONT_NAME
            title_style.font.size = Pt(26)
            title_style.font.bold = True
            title_style.paragraph_format.space_before = Pt(0)
            title_style.paragraph_format.space_after = Pt(12)
            title_style.paragraph_format.line_spacing = 1.0

        # Configure Heading styles (H1-H6)
        heading_sizes = {
            "Heading 1": Pt(20),
            "Heading 2": Pt(16),
            "Heading 3": Pt(14),
            "Heading 4": Pt(12),
            "Heading 5": Pt(11),
            "Heading 6": Pt(11),
        }
        heading_space_before = {
            "Heading 1": Pt(18),
            "Heading 2": Pt(14),
            "Heading 3": Pt(12),
            "Heading 4": Pt(10),
            "Heading 5": Pt(8),
            "Heading 6": Pt(8),
        }
        heading_space_after = {
            "Heading 1": Pt(8),
            "Heading 2": Pt(6),
            "Heading 3": Pt(6),
            "Heading 4": Pt(4),
            "Heading 5": Pt(4),
            "Heading 6": Pt(4),
        }

        for style_name, font_size in heading_sizes.items():
            if style_name in self.doc.styles:
                style = self.doc.styles[style_name]
                style.font.name = FONT_NAME
                style.font.size = font_size
                style.font.bold = True
                style.paragraph_format.space_before = heading_space_before.get(style_name, Pt(12))
                style.paragraph_format.space_after = heading_space_after.get(style_name, Pt(6))
                style.paragraph_format.line_spacing = 1.1

        # Configure List styles
        for list_style_name in ["List Bullet", "List Number", "List Paragraph"]:
            if list_style_name in self.doc.styles:
                list_style = self.doc.styles[list_style_name]
                list_style.font.name = FONT_NAME
                list_style.font.size = BODY_SIZE
                list_style.paragraph_format.space_before = Pt(2)
                list_style.paragraph_format.space_after = Pt(2)
                list_style.paragraph_format.line_spacing = LINE_SPACING

        # Configure Table styles
        if "Table Grid" in self.doc.styles:
            table_style = self.doc.styles["Table Grid"]
            # Table paragraph style settings
            table_style.font.name = FONT_NAME
            table_style.font.size = Pt(10)

    def write(
        self,
        result: OptimizationResult,
        output_path: Union[str, Path],
        source_url: Optional[str] = None,
        document_title: Optional[str] = None,
    ) -> Path:
        """
        Write the optimization result to a Word document.

        Args:
            result: OptimizationResult containing all optimized content.
            output_path: Path for the output .docx file.
            source_url: Optional source URL to display at top.
            document_title: Optional custom document title.

        Returns:
            Path to the created document.
        """
        output_path = Path(output_path)

        # Ensure .docx extension
        if output_path.suffix.lower() != ".docx":
            output_path = output_path.with_suffix(".docx")

        # Add document title if provided
        if document_title:
            title_para = self.doc.add_paragraph()
            title_para.style = "Title"
            title_para.add_run(document_title)

        # Add source URL if provided
        if source_url:
            url_para = self.doc.add_paragraph()
            url_para.add_run("Target Page: ")
            url_para.add_run(source_url)
            self.doc.add_paragraph()

        # Add reading guide
        self._add_reading_guide()

        # Add keyword plan table
        self._add_keyword_plan_table(
            primary_keyword=result.primary_keyword,
            secondary_keywords=result.secondary_keywords,
            keyword_usage_counts=result.keyword_usage_counts,
        )

        # Add meta elements table
        self._add_meta_table(result.meta_elements)

        # Add optimized content header
        self._add_section_header("OPTIMIZED CONTENT")

        # Add optimized blocks
        self._add_optimized_blocks(result.optimized_blocks)

        # Add FAQ section if present
        if result.faq_items:
            self._add_faq_section(result.faq_items)

        # Save document
        self.doc.save(str(output_path))

        return output_path

    def _add_reading_guide(self) -> None:
        """Add the reading guide section at the top."""
        # Add title
        title = self.doc.add_heading("Reading Guide", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Add guide text with highlighted sample using the marker protocol
        guide_para = self.doc.add_paragraph()
        guide_text = (
            "This document contains SEO-optimized content. Text highlighted in "
            "[[[ADD]]]green like this[[[ENDADD]]] indicates keyword insertions or "
            "SEO-focused adjustments. All non-highlighted text remains unchanged "
            "from the original content.\n\n"
            "Review the table below to see a summary of meta element changes, "
            "then scroll down to see the full optimized content."
        )
        add_marked_text(guide_para, guide_text)

        # Add spacing
        self.doc.add_paragraph()

    def _add_keyword_plan_table(
        self,
        primary_keyword: Optional[str],
        secondary_keywords: Optional[list[str]],
        keyword_usage_counts: Optional[dict[str, int]] = None,
    ) -> None:
        """Add the Keyword Plan table showing selected keywords and their usage counts.

        This table displays the primary keyword and secondary keywords
        that were used for optimization, along with how many times each
        keyword appears in the final optimized content.

        Args:
            primary_keyword: The primary keyword used for optimization.
            secondary_keywords: List of secondary keywords used.
            keyword_usage_counts: Dictionary mapping keyword phrase to occurrence count.
        """
        # Add section header
        self.doc.add_heading("Keyword Plan", level=2)

        # Ensure we have a usage counts dict
        if keyword_usage_counts is None:
            keyword_usage_counts = {}

        # Create table with 3 columns: Type, Keyword, and Usage Count
        table = self.doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.LEFT

        # Set column widths
        widths = [Inches(1.2), Inches(4.0), Inches(1.0)]
        for i, width in enumerate(widths):
            table.columns[i].width = width

        # Add header row
        header_cells = table.rows[0].cells
        headers = ["Type", "Keyword", "Usage Count"]
        for i, header in enumerate(headers):
            header_cells[i].text = header
            # Bold header text
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
            # Gray background for header
            set_cell_shading(header_cells[i], "D9D9D9")

        # Add primary keyword row
        if primary_keyword:
            row = table.add_row()
            cells = row.cells
            cells[0].text = "Primary"
            # Bold the "Primary" label
            for paragraph in cells[0].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
            cells[1].text = primary_keyword
            # Add usage count
            count = keyword_usage_counts.get(primary_keyword, 0)
            cells[2].text = str(count)

        # Add secondary keyword rows
        if secondary_keywords:
            for i, secondary in enumerate(secondary_keywords):
                row = table.add_row()
                cells = row.cells
                cells[0].text = f"Secondary {i + 1}"
                cells[1].text = secondary
                # Add usage count
                count = keyword_usage_counts.get(secondary, 0)
                cells[2].text = str(count)

        # Add spacing after table
        self.doc.add_paragraph()

    def _add_meta_table(self, meta_elements: list[MetaElement]) -> None:
        """Add the Current vs Optimized Meta Elements table."""
        # Add section header
        self.doc.add_heading("Current vs Optimized Meta Elements", level=2)

        # Create table with 4 columns
        table = self.doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.LEFT

        # Set column widths
        widths = [Inches(1.2), Inches(2.0), Inches(2.5), Inches(2.3)]
        for i, width in enumerate(widths):
            table.columns[i].width = width

        # Add header row
        header_cells = table.rows[0].cells
        headers = ["Element", "Current", "Optimized", "Why Changed"]
        for i, header in enumerate(headers):
            header_cells[i].text = header
            # Bold header text
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
            # Gray background for header
            set_cell_shading(header_cells[i], "D9D9D9")

        # Add data rows
        for meta in meta_elements:
            row = table.add_row()
            cells = row.cells

            # Element name
            cells[0].text = meta.element_name

            # Current value
            cells[1].text = meta.current or "(None)"

            # Optimized value - with highlights using add_marked_text
            cells[2].text = ""  # Clear existing content
            add_marked_text(cells[2].paragraphs[0], meta.optimized)

            # Why changed
            cells[3].text = meta.why_changed

        # Add spacing after table
        self.doc.add_paragraph()

    def _add_section_header(self, title: str) -> None:
        """Add a section header."""
        self.doc.add_heading(title, level=1)

    def _add_optimized_blocks(self, blocks: list[ParagraphBlock]) -> None:
        """Add optimized content blocks with highlights.

        Handles embedded newlines by splitting them into proper paragraphs/headings.
        """
        for block in blocks:
            if block.is_heading:
                # Add as heading with highlights
                level = min(block.heading_level.value, 6) if block.heading_level.value > 0 else 2
                heading = self.doc.add_heading(level=level)
                add_marked_text(heading, block.text)
            else:
                # Handle embedded newlines - split into proper paragraphs
                self._add_text_with_structure(block.text)

    def _add_text_with_structure(self, text: str) -> None:
        """
        Add text content, splitting on newlines into proper paragraphs/headings.

        Heuristically determines if a line is a heading based on:
        - Length (short lines without ending punctuation = likely headings)
        - Content patterns (bullet points, numbered lists)
        """
        # IMPORTANT: Normalize markers BEFORE splitting on newlines
        # This ensures markers don't get broken across lines
        text = _normalize_markers_for_newlines(text)

        # Final safety net: Fix "word.Word" patterns and spacing issues
        text = normalize_paragraph_spacing(text)

        if "\n" not in text:
            # No newlines - just add as single paragraph
            para = self.doc.add_paragraph()
            add_marked_text(para, text)
            return

        # Split on newlines and process each line
        lines = text.split("\n")

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # Detect if line is likely a heading
            is_heading = self._is_likely_heading(line)

            # Detect list items
            is_list_item = self._is_list_item(line)

            if is_heading:
                heading = self.doc.add_heading(level=2)
                add_marked_text(heading, line)
            elif is_list_item:
                # Add as a bullet point paragraph
                para = self.doc.add_paragraph(style="List Bullet")
                # Strip the leading bullet/dash/number
                clean_line = self._strip_list_prefix(line)
                add_marked_text(para, clean_line)
            else:
                para = self.doc.add_paragraph()
                add_marked_text(para, line)

    def _is_likely_heading(self, line: str) -> bool:
        """Determine if a line is likely a heading.

        Headings are typically:
        - Short (under ~80 chars)
        - Don't end with periods, commas, or other sentence-ending punctuation
        - Not a list item
        - Not starting with lowercase letter
        """
        # Skip list items
        if self._is_list_item(line):
            return False

        # Check length and punctuation
        stripped = line.strip()
        if not stripped:
            return False

        # Very long lines are not headings
        if len(stripped) > 80:
            return False

        # Lines ending with sentence punctuation are not headings
        if stripped.endswith((".", ",", ";", ":", "!", "?")):
            return False

        # Lines starting with lowercase are not headings (unless very short)
        if stripped[0].islower() and len(stripped) > 30:
            return False

        # Has some capital letters and looks title-like
        words = stripped.split()
        if len(words) <= 8:
            return True

        return False

    def _is_list_item(self, line: str) -> bool:
        """Check if line is a list item (bullet or numbered)."""
        stripped = line.strip()
        if not stripped:
            return False

        # Common bullet patterns
        if stripped.startswith(("-", "•", "*", "–", "—")):
            return True

        # Numbered list patterns (1., 1), a., a))
        if re.match(r"^\d+[\.\)]\s", stripped):
            return True
        if re.match(r"^[a-zA-Z][\.\)]\s", stripped):
            return True

        return False

    def _strip_list_prefix(self, line: str) -> str:
        """Remove list item prefix (bullet, dash, number)."""
        stripped = line.strip()

        # Remove common bullet prefixes
        for prefix in ["-", "•", "*", "–", "—"]:
            if stripped.startswith(prefix):
                return stripped[len(prefix):].strip()

        # Remove numbered prefixes
        match = re.match(r"^(?:\d+[\.\)]|[a-zA-Z][\.\)])\s*", stripped)
        if match:
            return stripped[match.end():].strip()

        return stripped

    def _add_faq_section(self, faq_items: list[FAQItem]) -> None:
        """Add the FAQ section."""
        # Add FAQ header
        faq_header = self.doc.add_heading("Frequently Asked Questions", level=2)
        # Mark the FAQ header as highlighted since it's new content
        for run in faq_header.runs:
            run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN

        for item in faq_items:
            # Question as bold paragraph (not heading) with highlights - all new content
            q_para = self.doc.add_paragraph()
            add_marked_text(q_para, item.question)
            # Make question text bold
            for run in q_para.runs:
                run.font.bold = True

            # Answer as paragraph with highlights - all new content
            a_para = self.doc.add_paragraph()
            add_marked_text(a_para, item.answer)

    def _add_content_table(self, table_block: TableBlock) -> None:
        """
        Add a content table to the document.

        Args:
            table_block: TableBlock containing rows and cells.
        """
        if not table_block.rows:
            return

        # Determine table dimensions
        num_rows = len(table_block.rows)
        num_cols = table_block.column_count

        if num_cols == 0:
            return

        # Create table
        table = self.doc.add_table(rows=num_rows, cols=num_cols)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.LEFT

        # Fill in cells
        for row_idx, row in enumerate(table_block.rows):
            for col_idx, cell in enumerate(row.cells):
                if col_idx < len(table.rows[row_idx].cells):
                    doc_cell = table.rows[row_idx].cells[col_idx]
                    # Clear default paragraph and add marked text
                    doc_cell.text = ""
                    add_marked_text(doc_cell.paragraphs[0], cell.text)

                    # Style header cells
                    if cell.is_header:
                        for paragraph in doc_cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
                        set_cell_shading(doc_cell, "D9D9D9")

        # Add spacing after table
        self.doc.add_paragraph()

    def _add_list_block(self, list_block: ListBlock) -> None:
        """
        Add a list to the document.

        Args:
            list_block: ListBlock containing list items.
        """
        style = "List Number" if list_block.ordered else "List Bullet"

        for item in list_block.items:
            para = self.doc.add_paragraph(style=style)
            add_marked_text(para, item.text)

    def _add_structured_block(self, block: ContentBlock) -> None:
        """
        Add a structured content block to the document.

        Args:
            block: ContentBlock to render.
        """
        if block.block_type in (ContentBlockType.PARAGRAPH, ContentBlockType.HEADING):
            if block.paragraph:
                if block.paragraph.is_heading:
                    level = min(block.paragraph.heading_level.value, 6)
                    if level > 0:
                        heading = self.doc.add_heading(level=level)
                        add_marked_text(heading, block.paragraph.text)
                else:
                    para = self.doc.add_paragraph()
                    add_marked_text(para, block.paragraph.text)

        elif block.block_type == ContentBlockType.TABLE:
            if block.table:
                self._add_content_table(block.table)

        elif block.block_type == ContentBlockType.LIST:
            if block.list_block:
                self._add_list_block(block.list_block)

    def add_structured_blocks(self, blocks: list[ContentBlock]) -> None:
        """
        Add a list of structured content blocks to the document.

        This method renders content with proper visual structure including
        headings, tables, lists, and paragraphs.

        Args:
            blocks: List of ContentBlock objects to render.
        """
        for block in blocks:
            self._add_structured_block(block)

    def _parse_marker_segments(self, text: str) -> list[tuple[str, bool]]:
        """
        Parse text with markers into segments.

        Returns list of (text, is_highlighted) tuples.

        Note: This method is kept for backwards compatibility with tests.
        The add_marked_text() function is the preferred way to add highlighted text.
        """
        segments: list[tuple[str, bool]] = []

        # Pattern to match [[[ADD]]]...[[[ENDADD]]]
        pattern = re.compile(
            rf"{re.escape(MARK_START)}(.*?){re.escape(MARK_END)}",
            re.DOTALL,
        )

        last_end = 0

        for match in pattern.finditer(text):
            # Add text before this match (not highlighted)
            if match.start() > last_end:
                before_text = text[last_end : match.start()]
                if before_text:
                    segments.append((before_text, False))

            # Add matched content (highlighted)
            highlighted_text = match.group(1)
            if highlighted_text:
                segments.append((highlighted_text, True))

            last_end = match.end()

        # Add any remaining text after last match
        if last_end < len(text):
            remaining = text[last_end:]
            if remaining:
                segments.append((remaining, False))

        # If no markers found, return original text unhighlighted
        if not segments:
            segments.append((text, False))

        return segments


def write_optimization_result(
    result: OptimizationResult,
    output_path: Union[str, Path],
) -> Path:
    """
    Convenience function to write optimization result to docx.

    Args:
        result: OptimizationResult to write.
        output_path: Output file path.

    Returns:
        Path to created document.
    """
    writer = DocxWriter()
    return writer.write(result, output_path)


def create_simple_docx_with_highlights(
    text: str,
    output_path: Union[str, Path],
    title: str = "Optimized Content",
) -> Path:
    """
    Create a simple docx with highlighted text.

    Useful for testing the highlight functionality.

    Args:
        text: Text with [[[ADD]]]...[[[ENDADD]]] markers.
        output_path: Output file path.
        title: Document title.

    Returns:
        Path to created document.
    """
    doc = Document()

    # Set Poppins font as default
    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Poppins"
    normal_style.font.size = Pt(11)
    normal_style.paragraph_format.space_before = Pt(6)
    normal_style.paragraph_format.space_after = Pt(6)
    normal_style.paragraph_format.line_spacing = 1.15

    # Configure Heading 1 style
    if "Heading 1" in doc.styles:
        h1_style = doc.styles["Heading 1"]
        h1_style.font.name = "Poppins"
        h1_style.font.size = Pt(20)
        h1_style.font.bold = True

    # Add title
    doc.add_heading(title, level=1)

    # Add content with highlights using the canonical add_marked_text function
    para = doc.add_paragraph()
    add_marked_text(para, text)

    output_path = Path(output_path)
    if output_path.suffix.lower() != ".docx":
        output_path = output_path.with_suffix(".docx")

    doc.save(str(output_path))
    return output_path
